from PyPDF2 import PdfReader
import docx

def parse_file(filepath):
    if filepath.lower().endswith(".pdf"):
        return parse_pdf(filepath)
    if filepath.lower().endswith(".docx"):
        return parse_docx(filepath)
    
    return None


def parse_docx(filepath):
    content = docx.Document(filepath)
    
    text = ""
    for i in range(len(content.paragraphs)):
        text += content.paragraphs[i].text + "\n" 
    
    return text

def parse_pdf(filepath):
    reader = PdfReader(filepath)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
        
    return text