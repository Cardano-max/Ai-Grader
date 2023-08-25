from docx import Document
from docx.shared import Mm
import os
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from uuid import uuid4
from docx.shared import Inches

def set_cell_margins(cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)


def add_single_row_data(document, title, value):
    paragraph = document.add_paragraph()
    paragraph.add_run(title).bold=True
    paragraph.add_run(value)

    return document


def add_image_insights(document, title, first_row, second_row):
    document.add_heading(title, 3)
    rows = document.add_table(rows=2, cols=4)
    
    for i, r in enumerate(first_row):
        cell1 = rows.cell(0, i)
        paragraph1 = cell1.paragraphs[0]
        paragraph1.add_run(r).bold = True

    for i, r in enumerate(second_row):
        cell1 = rows.cell(1, i)
        if r['type'] == "text":
            cell1.text = r['text']
        else:
            paragraph1 = cell1.paragraphs[0]
            run1 = paragraph1.add_run()
            run1.add_picture(r['path'], width=Inches(2), height=Inches(2))

    return document


def get_report(data):
    document = Document()
    # reduce left and right margin to 5mm (Optional)
    section = document.sections[0]
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)


    for row in [["Paritent Name: ", "Faizan Amin"], ["MRN: ", "11071981"], ["DOB: ", "November 7, 1981 (Age 51)"], ["Exam Date: ", datetime.now().strftime("%B %d, %Y")]]:
        document = add_single_row_data(document, row[0], row[1])

    document = add_image_insights(document, "Posterior Segment Imaging", 
                                ["OD", "Findings (Confidence)", "OS", "Findings (Confidence)"],
                                data)




    report_file = f"report_{uuid4()}.docx"
    document.save(report_file)
    return report_file
